from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from typing import List
import os
import asyncio
import pandas as pd
from PIL import Image

try:
    import pytesseract
    from pdf2image import convert_from_path
except ImportError:
    pytesseract = None
    convert_from_path = None


class DocumentProcessor:
    """
    Service for processing and chunking documents
    Supports PDF, scanned PDF OCR, images OCR, Word, Excel, TXT, Markdown, and CSV files
    """

    def __init__(
        self, chunk_size: int = 1000, chunk_overlap: int = 200
    ):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
        )

    async def process_document(
        self, file_path: str, file_type: str
    ) -> List[dict]:
        """
        Process a document and return chunks

        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, xlsx, xls, txt, md, csv)

        Returns:
            List of document chunks with metadata
        """
        # Load document based on type
        if file_type.lower() == "pdf":
            documents = await asyncio.to_thread(self._load_pdf, file_path)
            return await self._format_documents(file_path, documents)
        elif file_type.lower() in ["txt", "text", "md"]:
            loader = TextLoader(file_path)
        elif file_type.lower() == "csv":
            loader = CSVLoader(file_path)
        elif file_type.lower() == "docx":
            documents = await asyncio.to_thread(self._load_docx, file_path)
            return await self._format_documents(file_path, documents)
        elif file_type.lower() in ["xlsx", "xls"]:
            documents = await asyncio.to_thread(self._load_excel, file_path)
            return await self._format_documents(file_path, documents)
        elif file_type.lower() in ["png", "jpg", "jpeg", "webp"]:
            documents = await asyncio.to_thread(self._load_image_ocr, file_path)
            return await self._format_documents(file_path, documents)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        try:
            # CRITICAL FIX: Run synchronous Langchain tasks in a separate thread!
            documents = await asyncio.to_thread(loader.load)
            return await self._format_documents(file_path, documents)

        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")

    def _ensure_ocr_available(self) -> None:
        if pytesseract is None:
            raise ValueError("OCR is not installed. Please install pytesseract and tesseract.")

    def _load_pdf(self, file_path: str) -> List[Document]:
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        extracted_text = "\n".join(doc.page_content.strip() for doc in documents).strip()

        if len(extracted_text) >= 40:
            return documents

        return self._load_pdf_ocr(file_path)

    def _load_pdf_ocr(self, file_path: str) -> List[Document]:
        self._ensure_ocr_available()
        if convert_from_path is None:
            raise ValueError("PDF OCR is not installed. Please install pdf2image and poppler-utils.")

        pages = convert_from_path(file_path, dpi=200, first_page=1, last_page=10)
        documents = []
        for index, page in enumerate(pages, start=1):
            text = pytesseract.image_to_string(page).strip()
            if text:
                documents.append(
                    Document(
                        page_content=f"Page {index} OCR text:\n{text}",
                        metadata={"source": file_path, "page": index, "ocr": True},
                    )
                )

        if not documents:
            raise ValueError("OCR could not find readable text in this PDF.")

        return documents

    def _load_image_ocr(self, file_path: str) -> List[Document]:
        self._ensure_ocr_available()
        with Image.open(file_path) as image:
            text = pytesseract.image_to_string(image).strip()

        if not text:
            raise ValueError("OCR could not find readable text in this image.")

        return [Document(page_content=f"Image OCR text:\n{text}", metadata={"source": file_path, "ocr": True})]

    def _load_docx(self, file_path: str) -> List[Document]:
        doc = DocxDocument(file_path)
        parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table_index, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"Table {table_index}:\n" + "\n".join(rows))

        return [Document(page_content="\n\n".join(parts), metadata={"source": file_path})]

    def _load_excel(self, file_path: str) -> List[Document]:
        sheets = pd.read_excel(file_path, sheet_name=None)
        documents = []

        for sheet_name, frame in sheets.items():
            frame = frame.dropna(how="all").dropna(axis=1, how="all")
            if frame.empty:
                continue

            content = frame.to_csv(index=False)
            documents.append(
                Document(
                    page_content=f"Sheet: {sheet_name}\n{content}",
                    metadata={"source": file_path, "sheet": sheet_name},
                )
            )

        return documents

    async def _format_documents(self, file_path: str, documents: List[Document]) -> List[dict]:
        chunks = await asyncio.to_thread(self.text_splitter.split_documents, documents)

        formatted_chunks = []
        for i, chunk in enumerate(chunks):
            formatted_chunks.append(
                {
                    "text": chunk.page_content,
                    "metadata": {
                        **chunk.metadata,
                        "chunk_index": i,
                        "source_file": os.path.basename(file_path),
                    },
                }
            )

        return formatted_chunks

    def extract_text_from_chunks(self, chunks: List[dict]) -> str:
        """Extract plain text from chunks for context"""
        return "\n\n".join([chunk["text"] for chunk in chunks])
